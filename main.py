"""
TikTok Research AI Agent - FastAPI Backend
Stack: FastAPI + RapidAPI TikTok Scraper + Gemini + python-docx
"""

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional
import httpx
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import json
import os
import uuid
from datetime import datetime
from pathlib import Path

app = FastAPI(title="TikTok Research Agent")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- CONFIG ---
RAPIDAPI_KEY = os.getenv("RAPIDAPI_KEY")          # RapidAPI key
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")      # Google AI Studio key
genai.configure(api_key=GEMINI_API_KEY)

TIKWM_BASE = "https://tikwm.com/api"              # Free TikTok scraper API

# In-memory store (swap for SQLite/Postgres in production)
projects_db: dict = {}
analyses_db: dict = {}
briefs_db: dict = {}

Path("outputs").mkdir(exist_ok=True)


# ─────────────────────────────────────────
# MODELS
# ─────────────────────────────────────────

class Project(BaseModel):
    name: str
    client: str
    brand_bible: str          # Free-text: tone, audience, product, pillars
    brief_template: str       # Free-text template with {{placeholders}}

class SearchRequest(BaseModel):
    project_id: str
    keyword: str
    count: int = 10           # Number of videos to pull
    sort_type: int = 0        # 0=relevance, 1=likes, 2=date

class AnalyzeRequest(BaseModel):
    project_id: str
    video_ids: list[str]      # From search results

class BriefRequest(BaseModel):
    project_id: str
    analysis_ids: list[str]


# ─────────────────────────────────────────
# PROJECTS
# ─────────────────────────────────────────

@app.post("/projects")
async def create_project(project: Project):
    pid = str(uuid.uuid4())[:8]
    projects_db[pid] = {**project.dict(), "id": pid, "created_at": datetime.now().isoformat()}
    return projects_db[pid]

@app.get("/projects")
async def list_projects():
    return list(projects_db.values())

@app.get("/projects/{pid}")
async def get_project(pid: str):
    if pid not in projects_db:
        raise HTTPException(404, "Project not found")
    return projects_db[pid]

@app.put("/projects/{pid}")
async def update_project(pid: str, project: Project):
    if pid not in projects_db:
        raise HTTPException(404, "Project not found")
    projects_db[pid] = {**project.dict(), "id": pid, "created_at": projects_db[pid]["created_at"]}
    return projects_db[pid]


# ─────────────────────────────────────────
# TIKTOK SEARCH  (tikwm.com — free tier)
# ─────────────────────────────────────────

@app.post("/search")
async def search_tiktok(req: SearchRequest):
    """
    Search TikTok via tikwm.com API.
    Returns video metadata: id, desc, author, stats, cover_url, play_url
    """
    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.get(
            f"{TIKWM_BASE}/feed/search",
            params={
                "keywords": req.keyword,
                "count": req.count,
                "sort_type": req.sort_type,
                "cursor": 0,
            }
        )
    
    if resp.status_code != 200:
        raise HTTPException(502, f"TikTok API error: {resp.text}")
    
    data = resp.json()
    if data.get("code") != 0:
        raise HTTPException(502, f"TikTok API returned error: {data.get('msg')}")
    
    videos = []
    for v in data.get("data", {}).get("videos", []):
        videos.append({
            "id": v.get("video_id"),
            "desc": v.get("title"),
            "author": v.get("author", {}).get("nickname"),
            "author_handle": v.get("author", {}).get("unique_id"),
            "cover_url": v.get("cover"),
            "play_url": v.get("play"),           # Direct MP4 — Gemini can watch this
            "share_url": f"https://www.tiktok.com/@{v.get('author', {}).get('unique_id')}/video/{v.get('video_id')}",
            "likes": v.get("digg_count", 0),
            "comments": v.get("comment_count", 0),
            "shares": v.get("share_count", 0),
            "views": v.get("play_count", 0),
            "duration": v.get("duration", 0),
            "create_time": v.get("create_time"),
            "project_id": req.project_id,
        })
    
    return {"keyword": req.keyword, "videos": videos, "count": len(videos)}


# ─────────────────────────────────────────
# FETCH COMMENTS
# ─────────────────────────────────────────

async def fetch_comments(video_id: str, count: int = 30) -> list[str]:
    """Fetch top comments for a video via tikwm."""
    async with httpx.AsyncClient(timeout=20) as client:
        resp = await client.get(
            f"{TIKWM_BASE}/comment/list",
            params={"video_id": video_id, "count": count, "cursor": 0}
        )
    if resp.status_code != 200:
        return []
    data = resp.json()
    comments = []
    for c in data.get("data", {}).get("comments", []):
        text = c.get("text", "").strip()
        if text:
            comments.append(text)
    return comments[:count]


# ─────────────────────────────────────────
# GEMINI VIDEO ANALYSIS
# ─────────────────────────────────────────

HOOK_ANALYSIS_PROMPT = """
You are an expert creative strategist analyzing TikTok videos for a brand marketing team.

Analyze this TikTok video and provide a structured analysis in JSON format with these exact keys:

{
  "hook_type": "one of: Question | Bold Claim | Story | Shock/Surprise | Tutorial | POV | Challenge | Trend",
  "hook_text": "the exact opening line or text overlay if visible",
  "hook_score": 1-10 score for hook effectiveness,
  "emotional_angle": "the core emotion being triggered: curiosity / fear of missing out / aspiration / pain point / humor / etc",
  "format": "one of: Talking head | UGC demo | Voiceover B-roll | Duet | Text overlay / POV | Trend audio",
  "pacing": "Fast cut / Medium / Slow build",
  "visual_style": "brief description of visual aesthetic",
  "key_message": "core message/claim of the video in one sentence",
  "cta": "call to action if any",
  "why_it_works": "2-3 sentences on WHY this hook is effective for this audience",
  "swipe_risk": "High / Medium / Low — likelihood viewer swipes away in first 3 seconds",
  "inspiration_angles": ["3 content angle ideas inspired by this video that a brand could adapt"]
}

Return ONLY valid JSON, no markdown, no explanation.
"""

COMMENT_ANALYSIS_PROMPT = """
You are analyzing TikTok comments to extract audience intelligence for a creative strategist.

Here are the top comments from a viral TikTok video:

{comments}

Return a JSON object with these exact keys:
{
  "top_questions": ["list of top 3-5 questions the audience is asking"],
  "pain_points": ["list of pain points or frustrations mentioned"],
  "viral_phrases": ["memorable phrases or words the audience uses naturally"],
  "sentiment": "Positive / Mixed / Negative",
  "audience_signals": "2-3 sentences summarizing what this audience cares about and how they talk"
}

Return ONLY valid JSON.
"""


@app.post("/analyze")
async def analyze_videos(req: AnalyzeRequest):
    """
    For each video_id: fetch video data, run Gemini video analysis + comment analysis.
    Requires video data to be passed or re-fetched. In production, cache search results.
    """
    results = []
    model = genai.GenerativeModel("gemini-1.5-pro")

    for video_id in req.video_ids:
        try:
            # Fetch video info from tikwm
            async with httpx.AsyncClient(timeout=20) as client:
                resp = await client.get(f"{TIKWM_BASE}/", params={"url": f"https://www.tiktok.com/video/{video_id}"})
            
            vdata = resp.json().get("data", {})
            play_url = vdata.get("play") or vdata.get("wmplay", "")
            desc = vdata.get("title", "")
            
            # --- Gemini watches the video ---
            video_part = {"mime_type": "video/mp4", "file_uri": play_url} if play_url else None
            
            hook_analysis = {}
            if video_part:
                try:
                    gemini_resp = model.generate_content([
                        HOOK_ANALYSIS_PROMPT,
                        {"file_data": video_part}
                    ])
                    hook_analysis = json.loads(gemini_resp.text)
                except Exception as e:
                    # Fallback: analyze caption only
                    gemini_resp = model.generate_content([
                        f"Analyze this TikTok caption and return the JSON analysis:\n\nCaption: {desc}\n\n" + HOOK_ANALYSIS_PROMPT
                    ])
                    hook_analysis = json.loads(gemini_resp.text)

            # --- Fetch + analyze comments ---
            comments = await fetch_comments(video_id)
            comment_analysis = {}
            if comments:
                comments_text = "\n".join([f"- {c}" for c in comments])
                comment_resp = model.generate_content(
                    COMMENT_ANALYSIS_PROMPT.format(comments=comments_text)
                )
                comment_analysis = json.loads(comment_resp.text)

            analysis_id = str(uuid.uuid4())[:8]
            result = {
                "id": analysis_id,
                "video_id": video_id,
                "project_id": req.project_id,
                "desc": desc,
                "play_url": play_url,
                "cover_url": vdata.get("cover", ""),
                "likes": vdata.get("digg_count", 0),
                "views": vdata.get("play_count", 0),
                "hook_analysis": hook_analysis,
                "comment_analysis": comment_analysis,
                "analyzed_at": datetime.now().isoformat(),
            }
            analyses_db[analysis_id] = result
            results.append(result)

        except Exception as e:
            results.append({"video_id": video_id, "error": str(e)})

    return {"analyses": results}


@app.get("/analyses/{project_id}")
async def get_analyses(project_id: str):
    return [a for a in analyses_db.values() if a.get("project_id") == project_id]


# ─────────────────────────────────────────
# BRIEF GENERATION
# ─────────────────────────────────────────

BRIEF_GENERATION_PROMPT = """
You are a senior creative strategist at a top creative agency.

Generate a complete creative brief based on the following research.

## Brand Bible
{brand_bible}

## Brief Template
{brief_template}

## Research Insights (from {video_count} analyzed TikTok videos)
{research_summary}

Instructions:
- Fill every section of the brief template with specific, actionable content
- Use the brand bible to ensure everything is on-brand
- Pull specific hooks, phrases, and angles from the research
- Be concrete — include exact hook lines, not vague directions
- Write like a senior strategist who has watched every video and knows this audience cold

Return the completed brief as clean text, using the template structure exactly.
"""


@app.post("/generate-brief")
async def generate_brief(req: BriefRequest):
    project = projects_db.get(req.project_id)
    if not project:
        raise HTTPException(404, "Project not found")

    selected = [analyses_db[aid] for aid in req.analysis_ids if aid in analyses_db]
    if not selected:
        raise HTTPException(400, "No valid analyses found")

    # Build research summary
    summary_parts = []
    for a in selected:
        h = a.get("hook_analysis", {})
        c = a.get("comment_analysis", {})
        part = f"""
VIDEO: {a.get('desc', 'N/A')} | Views: {a.get('views', 0):,} | Likes: {a.get('likes', 0):,}
Hook Type: {h.get('hook_type')} | Hook: "{h.get('hook_text')}" | Score: {h.get('hook_score')}/10
Emotional Angle: {h.get('emotional_angle')} | Format: {h.get('format')}
Key Message: {h.get('key_message')}
Why It Works: {h.get('why_it_works')}
Inspiration Angles: {', '.join(h.get('inspiration_angles', []))}
Audience Questions: {', '.join(c.get('top_questions', []))}
Viral Phrases: {', '.join(c.get('viral_phrases', []))}
Pain Points: {', '.join(c.get('pain_points', []))}
"""
        summary_parts.append(part)

    research_summary = "\n---\n".join(summary_parts)

    model = genai.GenerativeModel("gemini-1.5-pro")
    resp = model.generate_content(
        BRIEF_GENERATION_PROMPT.format(
            brand_bible=project["brand_bible"],
            brief_template=project["brief_template"],
            video_count=len(selected),
            research_summary=research_summary,
        )
    )

    brief_text = resp.text
    brief_id = str(uuid.uuid4())[:8]

    # Save to DB
    brief_record = {
        "id": brief_id,
        "project_id": req.project_id,
        "project_name": project["name"],
        "client": project["client"],
        "content": brief_text,
        "analysis_ids": req.analysis_ids,
        "video_count": len(selected),
        "created_at": datetime.now().isoformat(),
    }
    briefs_db[brief_id] = brief_record

    return brief_record


@app.get("/briefs/{project_id}")
async def get_briefs(project_id: str):
    return [b for b in briefs_db.values() if b.get("project_id") == project_id]


# ─────────────────────────────────────────
# DOCX EXPORT
# ─────────────────────────────────────────

@app.get("/export/docx/{brief_id}")
async def export_docx(brief_id: str):
    from fastapi.responses import FileResponse

    brief = briefs_db.get(brief_id)
    if not brief:
        raise HTTPException(404, "Brief not found")

    doc = Document()

    # Title styling
    title = doc.add_heading(f"Creative Brief — {brief['client']}", 0)
    title.runs[0].font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)

    doc.add_paragraph(f"Project: {brief['project_name']}  |  Generated: {brief['created_at'][:10]}  |  Based on {brief['video_count']} videos")
    doc.add_paragraph("")

    # Write brief content — split on headers (lines starting with ##)
    for line in brief["content"].split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:])
        else:
            doc.add_paragraph(line)

    # Footer
    doc.add_paragraph("")
    footer_p = doc.add_paragraph(f"Generated by TikTok Research Agent · {datetime.now().strftime('%B %d, %Y')}")
    footer_p.runs[0].font.size = Pt(9)
    footer_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    filepath = f"outputs/brief_{brief_id}.docx"
    doc.save(filepath)

    return FileResponse(filepath, filename=f"brief_{brief['client']}_{brief_id}.docx",
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
